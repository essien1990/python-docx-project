from docx import Document
from docx.shared import Inches
import pyttsx3

engine = pyttsx3.init()
rate = engine.getProperty('rate')
engine.setProperty('rate', 250)
volume = engine.getProperty('volume')
engine.setProperty('volume', 2.0)


def speak(text):
    pyttsx3.speak(text)


document = Document()

document.add_picture('joe.jpg',
                     width=Inches(2.0)
                     )
name = input('what is your name? ')
speak('Hello ' + name + 'How are you today? ')

speak('What is your phone number? ')
phone_number = input('what is your phone number? ')
email = input('what is your email? ')


document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell about yourself? ')
)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ' '
)

p.add_run(experience_details)
# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')

    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To date ')
        # skills
        skills = input('Enter your skills ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + '  '
        )

    else:
        break


# skills
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')

    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p.document.add_paragraph(skill)
        p.style = 'list Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using AwesomeCode Course project using python and document library'

document.save('mycv.docx')
